from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

SCREENSHOTS = os.path.join(os.path.dirname(__file__), "screenshots")
OUTPUT = os.path.join(os.path.dirname(__file__), "Assignment_4_5_Report.docx")

doc = Document()

for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(3)
    section.right_margin  = Cm(2)

def font(run, size=11, bold=False, color=None):
    run.font.name = "Calibri"
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)

def h1(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after  = Pt(6)
    r = p.add_run(text)
    font(r, 16, True, (31, 73, 125))

def h2(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(4)
    r = p.add_run(text)
    font(r, 13, True, (31, 73, 125))

def h3(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(2)
    r = p.add_run(text)
    font(r, 11, True, (0, 0, 0))

def body(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(5)
    r = p.add_run(text)
    font(r, 11)

def note(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after  = Pt(5)
    p.paragraph_format.left_indent  = Cm(0.7)
    r = p.add_run(f"Note: {text}")
    font(r, 10, False, (80, 80, 80))

def bullet(text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(text)
    font(r, 11)

def table_row(cells, widths=None):
    return cells

def fig(filename, caption, width=15):
    path = os.path.join(SCREENSHOTS, filename)
    if os.path.exists(path):
        doc.add_picture(path, width=Cm(width))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cp = doc.add_paragraph()
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cp.paragraph_format.space_after = Pt(12)
        r = cp.add_run(f"Figure: {caption}")
        font(r, 9, False, (120, 120, 120))
    else:
        p = doc.add_paragraph()
        r = p.add_run(f"[Missing: {filename}]")
        font(r, 10, False, (200, 0, 0))

def pb():
    doc.add_page_break()

# ═══════════════════════════════════════════════
# TITLE PAGE
# ═══════════════════════════════════════════════
for _ in range(4):
    doc.add_paragraph()

tp = doc.add_paragraph()
tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = tp.add_run("ASSIGNMENT 4–5 REPORT")
font(r, 24, True, (31, 73, 125))

doc.add_paragraph()
tp2 = doc.add_paragraph()
tp2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = tp2.add_run("Incident Response Simulation and\nInfrastructure as Code Implementation")
font(r2, 14)

doc.add_paragraph()
tp3 = doc.add_paragraph()
tp3.alignment = WD_ALIGN_PARAGRAPH.CENTER
r3 = tp3.add_run("MicroShop — Containerized Microservices System")
font(r3, 12, True)

doc.add_paragraph()
tp4 = doc.add_paragraph()
tp4.alignment = WD_ALIGN_PARAGRAPH.CENTER
r4 = tp4.add_run("Introduction to SRE  |  2026")
font(r4, 11, False, (100, 100, 100))

doc.add_paragraph()
tp5 = doc.add_paragraph()
tp5.alignment = WD_ALIGN_PARAGRAPH.CENTER
r5 = tp5.add_run("GitHub: https://github.com/Nurlybek-Nurzhan/Assignment_4-5_SRE")
font(r5, 10, False, (31, 73, 125))

pb()

# ═══════════════════════════════════════════════
# SYSTEM OVERVIEW
# ═══════════════════════════════════════════════
h1("System Overview")
body(
    "MicroShop is a containerized e-commerce platform built to demonstrate SRE practices. "
    "It consists of five independent FastAPI microservices, a PostgreSQL database, an Nginx "
    "frontend acting as a reverse proxy, Prometheus for metrics collection with alert rules, "
    "and Grafana with auto-provisioned dashboards. All components run in Docker containers "
    "orchestrated by Docker Compose on a shared internal bridge network (microshop-net)."
)

h2("Architecture")
bullet("Auth Service (port 8001) — JWT authentication with bcrypt password hashing; credentials from environment variables")
bullet("Product Service (port 8002) — product catalog, persisted to a Docker volume via JSON file")
bullet("Order Service (port 8003) — order management, PostgreSQL backend; this is the incident target")
bullet("User Service (port 8004) — user profile management")
bullet("Chat Service (port 8005) — messaging with JWT token verification on all endpoints")
bullet("PostgreSQL — persistent database, used exclusively by the Order Service")
bullet("Prometheus (port 9090) — scrapes /metrics from all five services; alert rules defined for error detection")
bullet("Grafana (port 3000) — auto-provisioned datasource and dashboard on startup")
bullet("Nginx (port 80) — serves the frontend, proxies /api/* routes to backend services")

h2("Security Hardening Applied")
bullet("Auth credentials stored in .env file, never hardcoded in source code")
bullet("Passwords hashed with bcrypt — plaintext passwords never stored")
bullet("Real JWT tokens (HS256) issued on login; Chat Service validates token on every request")
bullet("All containers have memory and CPU resource limits (noisy-neighbor prevention)")
bullet("Docker healthchecks defined for all FastAPI services and Nginx")
bullet("terraform.tfstate excluded from git via .gitignore — contains sensitive AWS resource IDs")

h2("Known Limitations (Assignment Scope)")
bullet("Auth service uses two hardcoded users defined in .env — no user registration flow")
bullet("Chat messages are in-memory only — reset on container restart (no chat DB)")
bullet("Security group opens ports to 0.0.0.0/0 by default — the allowed_ssh_cidr variable allows restricting to a specific IP")
bullet("No SSH key pair configured for the EC2 instance — the ssh_key_name variable is available but left empty for this demo")

h2("Running System")
fig("1_terminal .png", "docker compose up — all 9 containers started successfully")
fig("2_docker_ps.png", "docker ps — all containers running with correct port mappings")

h2("Frontend Dashboard")
body(
    "The MicroShop Dashboard shows real-time health of all five services with green/red status dots. "
    "The header includes a login form — entering credentials calls the Auth Service and receives a JWT token "
    "which is then used for authenticated Chat Service calls."
)
fig("3_frontend_dashboard.png", "MicroShop Dashboard — all 5 services healthy, login form in header")

h2("Monitoring")
body(
    "Prometheus scrapes all five services at their /metrics endpoints every 15 seconds. "
    "Two alert rules are defined in alert_rules.yml: OrderServiceErrors (fires when order_errors_total "
    "increases) and ServiceDown (fires when any target becomes unreachable). "
    "Grafana loads the Prometheus datasource and the MicroShop dashboard automatically on startup "
    "via provisioning files — no manual configuration needed."
)
fig("4_prometheus_targets.png", "Prometheus — all 5 service targets showing UP")
fig("5_grafana.png", "Grafana — auto-provisioned dashboard with service health, request rates, and error panels")

pb()

# ═══════════════════════════════════════════════
# ASSIGNMENT 4 — INCIDENT REPORT
# ═══════════════════════════════════════════════
h1("ASSIGNMENT 4 — Incident Report")

h2("1. Incident Summary")
body(
    "On 2026-04-29, a production-level incident was simulated in the MicroShop Order Service "
    "by deliberately misconfiguring the DB_HOST environment variable from 'postgres' to 'wrong-host'. "
    "After restarting the container with this broken configuration, the Order Service lost all "
    "database connectivity and returned HTTP 503 on every request. All four other services "
    "remained fully operational throughout the incident, demonstrating correct fault isolation."
)

h2("2. Impact Assessment")
bullet("Order Service: 100% unavailable — /orders (GET and POST) and /health all returned HTTP 503")
bullet("Frontend dashboard: Order Service card showed 'Unreachable' with red status dot")
bullet("Users could not view existing orders or create new ones")
bullet("Auth, Product, User, and Chat services: fully operational — fault isolation confirmed")
bullet("Data integrity: no data was lost — PostgreSQL remained healthy throughout the incident")
bullet("Prometheus showed order-service as up=1 — the process was running but functionally broken")

h2("3. Severity Classification")
body("Severity: SEV-2 (Major) — a critical business function (order processing) was completely unavailable.")
body(
    "Classified SEV-2 rather than SEV-1 because the failure was isolated to one service and "
    "did not take down the entire platform. Authentication, product browsing, and other features "
    "remained functional."
)

h2("4. Timeline of Events")

tbl = doc.add_table(rows=1, cols=2)
tbl.style = "Table Grid"
hdr = tbl.rows[0].cells
hdr[0].text = "Time"
hdr[1].text = "Event"
for cell in hdr:
    for p in cell.paragraphs:
        for r in p.runs:
            font(r, 10, True)

events = [
    ("T+00:00", "DB_HOST changed from 'postgres' to 'wrong-host' in docker-compose.yml"),
    ("T+00:05", "Order Service restarted: docker compose up -d --force-recreate order-service"),
    ("T+00:10", "Frontend dashboard detected Order Service as 'Unreachable'"),
    ("T+00:15", "/api/orders/orders returned HTTP 503 Service Unavailable"),
    ("T+00:20", "order_errors_total counter began incrementing in Prometheus"),
    ("T+00:25", "/health returned 503: 'DB connection failed: could not translate host name wrong-host'"),
    ("T+00:30", "docker logs confirmed repeated psycopg2 DNS resolution failures"),
    ("T+01:00", "Root cause confirmed: incorrect DB_HOST environment variable"),
    ("T+01:05", "DB_HOST corrected back to 'postgres' in docker-compose.yml"),
    ("T+01:10", "Order Service restarted with corrected configuration"),
    ("T+01:15", "/health returned HTTP 200: {status: ok, db: connected}"),
    ("T+01:20", "Frontend dashboard confirmed all services green — full recovery"),
]
for t, e in events:
    row = tbl.add_row().cells
    row[0].text = t
    row[1].text = e
    for cell in row:
        for p in cell.paragraphs:
            for r in p.runs:
                font(r, 10)

doc.add_paragraph()

h2("5. Root Cause Analysis")
body(
    "Root cause: misconfigured DB_HOST environment variable. "
    "Docker Compose services resolve each other by container name through Docker's internal DNS. "
    "The hostname 'postgres' resolves correctly on the microshop-net network because a container "
    "with that name exists. The value 'wrong-host' has no DNS record on the network, so every "
    "call to psycopg2.connect() immediately raised an OperationalError at the DNS resolution step "
    "— before even attempting a TCP connection to port 5432."
)
body(
    "Key SRE observation: Prometheus showed order-service as up=1 throughout the entire incident. "
    "This is because Prometheus scrapes /metrics, which does not touch the database and always "
    "responds successfully. The service process was alive, but functionally broken. "
    "This demonstrates a critical distinction in SRE: process liveness does not equal service health. "
    "The /health endpoint (which calls psycopg2.connect() on every request) correctly returned 503 "
    "and provided the exact error message pointing to the bad hostname."
)

h2("Incident Evidence")
fig("6_incident_dashboard.png",
    "Dashboard during incident — Order Service 'Unreachable', 503 errors visible in Network tab")
fig("7_incident_prometheus.png",
    "Prometheus during incident — up=1 for all services including order-service. "
    "The process is alive but /metrics does not check DB. The real signal is order_errors_total.")
fig("8a_incident_logs.png",
    "Docker logs — alternating pattern: GET /orders → 503, GET /metrics → 200. "
    "Confirms service is running but every DB request fails.")
fig("8b_incident_metrics_and_health.png",
    "/health returning 503 with exact error: 'could not translate host name wrong-host' — "
    "this is the root cause evidence. Right panel shows order_errors_total incrementing.")

h2("6. Mitigation Steps")
bullet("Step 1 — Identified failing container: docker logs order-service showed repeated 503 on /orders")
bullet("Step 2 — Checked /health endpoint: returned exact error message with the bad hostname 'wrong-host'")
bullet("Step 3 — Inspected docker-compose.yml: found DB_HOST set to 'wrong-host' instead of 'postgres'")
bullet("Step 4 — Applied fix: changed DB_HOST back to 'postgres'")
bullet("Step 5 — Restarted only the affected container: docker compose up -d --force-recreate order-service")
bullet("Step 6 — Verified recovery: /health returned 200 with db: connected")
bullet("Step 7 — Confirmed via dashboard: all services showing green")

h2("7. Resolution Confirmation")
body(
    "Full recovery was achieved within approximately 70 seconds of applying the fix. "
    "The /health endpoint returned HTTP 200 with db: connected, the frontend dashboard showed "
    "all five services green, and the order_errors_total counter stopped incrementing in Prometheus. "
    "No data was lost and no other services required any intervention."
)
fig("9_recovery_dashboard.png",
    "Recovery — all services green, Order Service showing DB: connected")
fig("10a_recovery_logs.png",
    "Recovery logs — HTTP 200 responses on /orders confirmed after fix")
fig("10b_recovery_metrics_and_health.png",
    "/health returning 200 with db: connected after DB_HOST corrected")

pb()

# ═══════════════════════════════════════════════
# POSTMORTEM
# ═══════════════════════════════════════════════
h1("Postmortem Analysis")

h2("1. Incident Overview")
body(
    "On 2026-04-29, the MicroShop Order Service became fully unavailable for approximately 70 seconds "
    "due to an incorrect database hostname in the service environment configuration. The incident was "
    "intentionally introduced as part of Assignment 4 to simulate a realistic production failure and "
    "practice the complete incident response lifecycle — detection, analysis, mitigation, and recovery."
)

h2("2. Customer Impact")
bullet("100% of order-related requests (view orders, create order) failed with HTTP 503")
bullet("The MicroShop frontend showed Order Service as 'Unreachable' for the duration of the incident")
bullet("No data loss occurred — PostgreSQL remained healthy and all existing orders were preserved")
bullet("Product browsing, authentication, user profiles, and chat were fully unaffected")
bullet("In a real production environment this would represent complete loss of revenue-generating functionality")

h2("3. Root Cause Analysis")
body(
    "The DB_HOST environment variable was manually changed to a non-existent hostname. "
    "In Docker Compose, services discover each other by container name via Docker's internal DNS. "
    "Setting DB_HOST='wrong-host' caused immediate DNS lookup failure on every database operation. "
    "The service had no connection retry logic — each incoming request attempted a fresh "
    "psycopg2.connect() with a 3-second timeout, failed instantly, incremented order_errors_total, "
    "and returned HTTP 503."
)
body(
    "Contributing factor: Prometheus monitoring showed the service as UP throughout the incident "
    "because it only checks /metrics endpoint availability (process liveness), not actual DB "
    "connectivity (functional health). This created a misleading signal that could delay detection "
    "in a real environment without proper alerting on error rates."
)

h2("4. Detection and Response Evaluation")
body(
    "Detection was fast — the frontend dashboard showed the failure within 10 seconds of the "
    "container restart. The /health endpoint was the critical diagnostic tool: it provided the "
    "exact error message including the bad hostname, making root cause identification immediate."
)
body(
    "Gap identified: Prometheus showed up=1 during the incident. In production, relying solely on "
    "process liveness metrics would delay detection. The OrderServiceErrors alert rule added to "
    "alert_rules.yml addresses this — it fires when order_errors_total increases, catching "
    "functional failures that process liveness cannot detect."
)

h2("5. Resolution Summary")
body(
    "A single configuration change (correcting DB_HOST to 'postgres') followed by a container "
    "restart fully restored the service in under 2 minutes. No code changes, no data recovery, "
    "and no other service required any intervention. The microservices fault isolation architecture "
    "worked exactly as designed."
)

h2("6. Lessons Learned")
bullet("Environment variables are a frequent source of misconfiguration — they should be validated at service startup")
bullet("Process liveness (up=1 in Prometheus) does not mean a service is functionally healthy")
bullet("Dependency health checks in /health endpoints are essential — they reveal failures that Prometheus scraping misses")
bullet("Microservices fault isolation worked correctly — a broken Order Service did not cascade to other services")
bullet("The /health endpoint design (checking actual DB connectivity on every call) was the correct approach and proved essential")
bullet("Manual changes to running configurations should always be validated before applying in production")

h2("7. Action Items")
bullet("DONE — Added Prometheus alert rules: OrderServiceErrors and ServiceDown in monitoring/alert_rules.yml")
bullet("DONE — Grafana dashboard auto-provisioned with error rate panels visible during incident")
bullet("DONE — Resource limits and Docker healthchecks added to all containers")
bullet("TODO — Add DB_HOST validation at Order Service startup: attempt connection and fail fast with clear error")
bullet("TODO — Implement connection retry with exponential backoff in psycopg2 calls")
bullet("TODO — Configure Grafana alert notification channel (email/Slack) connected to Prometheus alerts")
bullet("TODO — Replace 0.0.0.0/0 in Terraform security group with specific IP ranges in production")

pb()

# ═══════════════════════════════════════════════
# ASSIGNMENT 5 — TERRAFORM
# ═══════════════════════════════════════════════
h1("ASSIGNMENT 5 — Infrastructure as Code (Terraform + AWS)")

h2("Overview")
body(
    "The cloud infrastructure for MicroShop was provisioned using Terraform with the AWS provider. "
    "Terraform is a declarative IaC tool — the entire infrastructure is described in .tf files "
    "and created, updated, or destroyed with single commands. This makes infrastructure reproducible, "
    "version-controlled, and auditable — unlike manual AWS Console operations."
)

h2("Implementation: Two Phases")
body(
    "During development, Terraform was first configured with a local Docker provider "
    "(kreuzwerker/docker) to validate the configuration structure locally. "
    "This phase is visible in screenshots 13_terraform_apply_*. "
    "After the professor confirmed that cloud deployment was required, the configuration "
    "was migrated to the AWS provider. The final AWS deployment is shown in screenshots "
    "14_terraform_apply_after_AWS_implementation_*. "
    "The code in the repository reflects the final AWS configuration only."
)
note(
    "The AWS apply created: EC2 instance at 34.238.137.187 (us-east-1), "
    "security group sg-0091acf407323802f. The Docker provider files were removed from state "
    "before the AWS migration."
)

h2("Terraform Files")

h3("main.tf")
body(
    "Defines two AWS resources. aws_security_group creates a firewall with inbound rules for "
    "SSH (22), HTTP (80), Grafana (3000), and Prometheus (9090). SSH and monitoring ports use "
    "the allowed_ssh_cidr variable — set it to your IP in production to restrict access. "
    "aws_instance provisions a t3.micro EC2 VM with Amazon Linux 2. The user_data script "
    "runs on first boot and installs Docker, git, and docker-compose so the application "
    "can be deployed immediately after the instance starts. "
    "The key_name field accepts an AWS key pair name via the ssh_key_name variable — "
    "left empty for this demo but required for SSH access in a real deployment."
)

h3("variables.tf")
body(
    "Declares five input variables: aws_region, instance_type, ami_id, allowed_ssh_cidr, "
    "and ssh_key_name. Using variables instead of hardcoded values makes the configuration "
    "reusable — changing the region, instance size, or IP restrictions requires editing "
    "only terraform.tfvars, not the core infrastructure code."
)

h3("terraform.tfvars")
body(
    "Provides actual values: us-east-1, t3.micro, Amazon Linux 2 AMI. "
    "t3.micro was chosen because t2.micro was rejected by AWS as not free-tier eligible "
    "in the selected region. Includes comments explaining how to restrict access in production."
)

h3("outputs.tf")
body(
    "Defines what Terraform prints after apply: instance_public_ip, instance_public_dns, "
    "grafana_url, prometheus_url, and security_group_id. These make it easy to access "
    "deployed services without looking up the IP manually in the AWS Console."
)

h2("Security Considerations")
bullet("SSH and monitoring ports use allowed_ssh_cidr variable — defaults to 0.0.0.0/0 for demo, restrict to specific IP in production")
bullet("No SSH key pair configured by default — set ssh_key_name variable to an existing AWS key pair to enable SSH access")
bullet("Grafana and Prometheus have no built-in auth when exposed publicly — in production place behind a VPN or add authentication")
bullet("terraform.tfstate is excluded from git via .gitignore — it contains AWS resource IDs and should never be committed")

h2("Deployment Commands")
bullet("terraform init — downloads hashicorp/aws provider plugin, initializes working directory")
bullet("terraform plan — previews all changes; shows exact resources to be created before any AWS calls")
bullet("terraform apply — creates resources on AWS (completed in ~30 seconds)")
bullet("terraform destroy — removes all resources cleanly; always run after assignment to avoid AWS charges")

h2("Terraform Screenshots")
fig("11a_terraform_init.png",
    "terraform init — AWS provider (hashicorp/aws v5.x) downloaded successfully")
fig("11b_aws_credentials.png",
    "AWS IAM credentials configured via access key for Terraform authentication")
fig("12_terraform_plan_0.png",
    "terraform plan — showing aws_security_group and aws_instance to be created")
fig("12_terraform_plan_1.png",
    "terraform plan — EC2 instance details: AMI, instance type, security group, user_data")
fig("14_terraform_apply_after_AWS_implementation_0.png",
    "terraform apply (AWS) — EC2 instance creation in progress on AWS")
fig("14_terraform_apply_after_AWS_implementation_1.png",
    "terraform apply (AWS) — security group sg-0091acf407323802f created")
fig("14_terraform_apply_after_AWS_implementation_2.png",
    "terraform apply (AWS) — outputs: public IP 34.238.137.187, Grafana and Prometheus URLs")

h2("Why Terraform over Manual AWS Console")
bullet("Reproducible — running terraform apply again produces identical infrastructure every time")
bullet("Version controlled — infrastructure changes tracked in git alongside application code")
bullet("Self-documenting — .tf files describe exactly what exists and why")
bullet("Automated — no manual clicking, no missed steps, no human error")
bullet("Destroyable — terraform destroy cleanly removes everything, preventing surprise AWS charges")

pb()

# ═══════════════════════════════════════════════
# DEPLOYMENT GUIDE
# ═══════════════════════════════════════════════
h1("Deployment Guide")

h2("Prerequisites")
bullet("Docker Desktop installed and running")
bullet("Terraform >= 1.0 installed")
bullet("AWS account with IAM user (AdministratorAccess policy, access key + secret key)")
bullet("Git")

h2("1. Clone the Repository")
body("git clone https://github.com/Nurlybek-Nurzhan/Assignment_4-5_SRE.git")
body("cd Assignment_4-5_SRE")

h2("2. Configure Environment")
body("cp .env.example .env")
body("Edit .env and set your preferred passwords and JWT secret.")

h2("3. Run Locally with Docker Compose")
body("docker compose up --build")
body("Dashboard: http://localhost:80  |  Prometheus: http://localhost:9090  |  Grafana: http://localhost:3000")
body("Login with credentials from .env (default: nurzhan / password123).")

h2("4. Provision AWS Infrastructure")
body("cd terraform")
body("aws configure   # enter IAM access key, secret key, region: us-east-1")
body("terraform init && terraform plan && terraform apply")
body("SSH into the instance, clone the repo, and run docker compose up.")

h2("5. Simulate the Incident")
body("In docker-compose.yml change DB_HOST: postgres  →  DB_HOST: wrong-host")
body("docker compose up -d --force-recreate order-service")
body("Observe Order Service become Unreachable. Check /health for the error message. Watch order_errors_total in Prometheus.")

h2("6. Recover from the Incident")
body("Restore DB_HOST: postgres in docker-compose.yml")
body("docker compose up -d --force-recreate order-service")
body("All services return to healthy state within ~10 seconds.")

h2("7. Clean Up AWS Resources")
body("cd terraform && terraform destroy")
body("Always run this after the assignment to avoid ongoing AWS charges.")

# ═══════════════════════════════════════════════
# SAVE
# ═══════════════════════════════════════════════
doc.save(OUTPUT)
print(f"Saved: {OUTPUT}")
